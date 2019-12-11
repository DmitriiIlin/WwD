from docx import Document
from openpyxl import Workbook, load_workbook
import random, os
import PyPDF2


def docx_creater (word_name, inputtext_1, inputtext_2):
    #Создаем документ docx
    doc = Document()
    doc.add_heading(inputtext_1, 0)
    p = doc.add_paragraph(inputtext_2)
    doc.save(str(word_name))

def excel_creater (excel_file_name, sheet_name):
    #Создаем excel документ
    wb = Workbook()
    start_column = 1
    start_row = 1
    work_sheet = wb.create_sheet(title=sheet_name)
    for i in range(0,12):
        for j in range(0,12):
            work_sheet.cell(column = start_column+i, row = start_row+j, value = random.randint(0,100) )
    wb.save(filename = excel_file_name)

def open_pdf (path_to_file, file_name):
    #Пробуем вывести информацию из pdf файла
    current_working_directory=os.getcwd()
    os.chdir(path_to_file)
    pdf_file=open(file_name,'rb')
    read_pdf=PyPDF2.PdfFileReader(pdf_file)
    page=read_pdf.getPage(1)
    page_text=page.extractText()
    print(page_text.encode("utf-8"))
    os.chdir(current_working_directory)


